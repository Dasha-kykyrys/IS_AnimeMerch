from PyQt5 import QtWidgets, QtCore
from PyQt5.QtCore import Qt, QSortFilterProxyModel, QRegExp, QDate
from PyQt5.QtGui import QStandardItemModel, QStandardItem
import sys
from PyQt5.QtWidgets import QFileDialog
from GUI import Ui_MainWindow, Ui_CreateEditAnime_form, Ui_CreateEditProduct_form, Ui_preview_form, ItemDelegateData, ItemDelegateCheck
from database import load_data_from_db, add_to_database_product, add_to_database_anime, delete_anime_by_name, \
    delete_product, update_anime, update_product, save_database, add_sales
import re
from docx import Document
from datetime import datetime
import os
from docx2pdf import convert

class FilterOrder(QSortFilterProxyModel):
    def __init__(self, *args, start_date=None, end_date=None, **kwargs):
        super().__init__(*args, **kwargs)
        self.start_date = start_date
        self.end_date = end_date

    def set_date_range(self, start_date, end_date):
        self.start_date = start_date
        self.end_date = end_date
        self.invalidateFilter()

    def filterAcceptsRow(self, source_row, source_parent):
        model = self.sourceModel()

        # Дата продажи из модели
        sale_date_index = model.index(source_row, 5, source_parent)
        sale_date = model.data(sale_date_index)
        sale_date = sale_date.split()[0]

        if isinstance(sale_date, str):
            sale_date = QDate.fromString(sale_date, 'yyyy-MM-dd')

        if self.start_date and self.end_date:
            return self.start_date <= sale_date <= self.end_date

        return True  # Если даты не заданы

class FilterCatalog(QSortFilterProxyModel):
    def filterAcceptsRow(self, source_row, source_parent):
        model = self.sourceModel()
        filter_text = self.filterRegExp().pattern().lower()

        if not filter_text:
            return True

        for column in [1, 2, 3]:  # Индексы для Наименование, Аниме и Цена
            index = model.index(source_row, column, source_parent)
            if filter_text in str(model.data(index)).lower():
                return True

        return False

class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(MainWindow, self).__init__()
        self.ui_main_window = Ui_MainWindow()
        self.ui_main_window.setupUi(self)

        #Таблицы во вкладке управления
        self.ui_main_window.animetable_button.clicked.connect(self.open_anime_table)
        self.ui_main_window.productstable_button.clicked.connect(self.open_product_table)

        #открытие окон
        self.ui_main_window.print_sales_button.clicked.connect(self.open_preview)
        self.ui_main_window.print_button.clicked.connect(self.open_preview)
        self.ui_main_window.creat_button.clicked.connect(self.open_create)

        # Модель таблицы Каталог
        self.catalog_model = QStandardItemModel(0, 5, self)  # 5 столбцов
        self.set_model_catalog()

        # Подключение кнопки поиска в каталоге
        self.ui_main_window.find_button.clicked.connect(self.on_search_clicked)

        # Подключение кнопки фильтрации в продажах
        self.ui_main_window.confirmdate_button.clicked.connect(self.on_filter_clicked)
        self.ui_main_window.begin_dateEdit.setDate(datetime.now())
        self.ui_main_window.end_dateEdit.setDate(datetime.now())

        # Смена вкладок
        self.ui_main_window.catalog_button.clicked.connect(lambda: self.change_widget(0))
        self.ui_main_window.sales_button.clicked.connect(lambda: self.change_widget(1))
        self.ui_main_window.management_button.clicked.connect(lambda: self.change_widget(2))

        # Модель таблицы Чек
        self.check_model = QStandardItemModel(0, 3, self)
        self.set_model_check()

        # Модель таблицы Продажа
        self.sale_model = QStandardItemModel(0, 5, self)
        self.set_model_sale()

        # Модель таблицы Товар
        self.product_model = QStandardItemModel(0, 5, self)
        self.open_product_table()  # Таблица товары

        # Модель таблицы Аниме
        self.anime_model = QStandardItemModel(0, 2, self)

        # Окно создания/редактирования аниме/товара
        self.create_edit_product = CreateEditProduct(self.ui_main_window, self.product_model, self.add_delegate_management)
        self.create_edit_anime = CreateEditAnime(self.ui_main_window, self.anime_model, self.add_delegate_management)

        # Словарь для хранения информации о ценах и количестве
        self.item_info = {}
        self.total_cost = 0
        self.ui_main_window.total_label.setText(str(self.total_cost) + " руб.")

        self.ui_main_window.add_button.clicked.connect(self.add_item_to_check) # Кнопка добавление товара в чек
        self.ui_main_window.remove_button.clicked.connect(self.remove_item_from_check) # Кнопка удаления товара из чека

        self.ui_main_window.confirm_button.clicked.connect(self.save_check) # Кнопка создания файла чека

    # Настройки для модели таблицы "Чек"
    def set_model_check(self):
        self.check_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Аниме', 'Кол-во'])
        self.ui_main_window.chek_table.setModel(self.check_model)
        for column in range(self.check_model.columnCount()):  # Адаптирующиеся под данные столбцы
            self.ui_main_window.chek_table.resizeColumnToContents(column)

    # Настройки для модели таблицы "Продажа"
    def set_model_sale(self):
        self.proxy_model_sale = FilterOrder()

        load_data_from_db(self.sale_model, 'sale_table')
        self.sale_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Аниме', 'Цена', 'Кол-во', 'Дата'])

        self.proxy_model_sale.setSourceModel(self.sale_model)
        self.ui_main_window.sales_table.setModel(self.proxy_model_sale)
        for column in range(self.sale_model.columnCount()):  # Адаптирующиеся под данные столбцы
            self.ui_main_window.sales_table.resizeColumnToContents(column)

    # Настройки для модели таблицы "Товар"
    def set_model_product(self):
        load_data_from_db(self.product_model, 'product_table')
        self.product_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Аниме', 'Цена', 'Кол-во', 'Действие'])
        self.ui_main_window.management_table.setModel(self.product_model)
        self.add_delegate_management(self.product_model, self.ui_main_window.management_table, 5)
        for column in range(self.product_model.columnCount()): # Адаптирующиеся под данные столбцы
            self.ui_main_window.management_table.resizeColumnToContents(column)

    # Настройки для модели таблицы "Каталог"
    def set_model_catalog(self):
        self.proxy_model_catalog = FilterCatalog()
        load_data_from_db(self.catalog_model, 'product_table')
        self.catalog_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Аниме', 'Цена', 'Кол-во'])

        self.proxy_model_catalog.setSourceModel(self.catalog_model)
        self.ui_main_window.catalog_table.setModel(self.proxy_model_catalog)
        for column in range(self.catalog_model.columnCount()):  # Адаптирующиеся под данные столбцы
            self.ui_main_window.catalog_table.resizeColumnToContents(column)

    # Открытие таблицы "Аниме"
    def open_anime_table(self):
        self.ui_main_window.print_button.hide()
        self.ui_main_window.productstable_button.setEnabled(True)
        self.ui_main_window.animetable_button.setEnabled(False)

        load_data_from_db(self.anime_model, 'anime_table')
        self.anime_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Действие'])
        self.ui_main_window.management_table.setModel(self.anime_model)
        self.add_delegate_management(self.anime_model, self.ui_main_window.management_table, 2)
        for column in range(self.anime_model.columnCount()): # Адаптирующиеся под данные столбцы
            self.ui_main_window.management_table.resizeColumnToContents(column)

    # Открытие таблицы "Товар"
    def open_product_table(self):
        self.ui_main_window.print_button.show()
        self.ui_main_window.animetable_button.setEnabled(True)
        self.ui_main_window.productstable_button.setEnabled(False)

        self.set_model_product()

    def on_filter_clicked(self):
        start_date = self.ui_main_window.begin_dateEdit.date()  # Начальная дата
        end_date = self.ui_main_window.end_dateEdit.date()  # Конечная дата

        # Диапазон дат в прокси модели
        self.proxy_model_sale.set_date_range(start_date, end_date)

    # Функция поиска по символу в таблицы каталог
    def on_search_clicked(self):
        # текст из поля поиска
        text = self.ui_main_window.search_field.toPlainText()
        self.proxy_model_catalog.setFilterRegExp(QRegExp(text, Qt.CaseInsensitive, QRegExp.RegExp))

    # Добавление кнопок в столбик кол-во в таблице "Чек"
    def add_delegate_check(self, model, table, number):
        for row in range(model.rowCount()):
            item = model.item(row, 1)

            if item is not None:
                item_delegate = ItemDelegateCheck(row)  # экземпляр ItemDelegateCheck
                table.setIndexWidget(model.index(row, number), item_delegate)  # Установка кнопок в таблицу
                item_delegate.button_clicked.connect(self.handle_button_click_check)

                # Чтобы не сбрасывалось уже набранное количество товара в чеке
                item_number = self.check_model.item(row, 0).text()
                item_info = self.item_info.get(item_number)
                if item_info is None:
                    return  # Если информация не найдена
                current_count = int(item_info['select_quantity'])
                item_delegate.update_count_label(current_count)

    def handle_button_click_check(self, button_type, row_number):
        type_button = button_type  # Определение типа кнопки
        number_row = row_number  # Определение строки
        index_count = self.ui_main_window.chek_table.model().index(number_row, 3)
        item_delegate = self.ui_main_window.chek_table.indexWidget(index_count)

        # Получение данных из выбранной строки
        item_number = self.check_model.item(number_row, 0).text()  # №

        item_info = self.item_info.get(item_number)
        if item_info is None:
            return  # Если информация не найдена

        item_quantity = int(item_info['quantity'])
        item_select_quantity = int(item_info['select_quantity'])
        item_price = int(item_info['price'])
        item_delegate.update_count_label(item_select_quantity)

        if type_button == "add" and item_select_quantity < item_quantity:
            item_select_quantity = item_select_quantity + 1  # Увеличиваем счетчик
            item_delegate.update_count_label(item_select_quantity)
            self.total_cost += int(item_price)
            self.ui_main_window.total_label.setText(str(self.total_cost) + " руб.")  # Увеличение суммы покупки

        elif type_button == "delete" and item_select_quantity > 1:
            item_select_quantity = item_select_quantity - 1  # Уменьшаем счетчик
            item_delegate.update_count_label(item_select_quantity)
            self.total_cost -= int(item_price)
            self.ui_main_window.total_label.setText(str(self.total_cost) + " руб.")  # Уменьшение суммы покупки

        item_info['select_quantity'] = str(item_select_quantity)

    # Функция добавления позиций в чек
    def add_item_to_check(self):
        # Получение индексы строки в таблице каталога
        selected_indexes = self.ui_main_window.catalog_table.selectedIndexes()

        if not selected_indexes:
            return  # Если ничего не выбрано

        # Получение индекса строки, выбранной пользователем в прокси-модели
        proxy_index = selected_indexes[0]
        # Преобразование индекса прокси-модели в индекс исходной модели
        source_index = self.proxy_model_catalog.mapToSource(proxy_index)

        # Получение данных из исходной модели
        item_number = self.catalog_model.item(source_index.row(), 0).text()  # №
        item_name = self.catalog_model.item(source_index.row(), 1).text()  # Наименование
        item_anime = self.catalog_model.item(source_index.row(), 2).text()  # Аниме
        item_price = self.catalog_model.item(source_index.row(), 3).text()  # Цена
        item_quantity = self.catalog_model.item(source_index.row(), 4).text()  # Кол-во
        item_select_quantity = 1

        # Сохранение информации о цене и количестве товара в словарь
        self.item_info[item_number] = {
            'name': item_name,
            'anime': item_anime,
            'price': item_price,
            'quantity': item_quantity,
            'select_quantity': item_select_quantity,
        }

        # Добавление элемента в чек
        new_row = self.check_model.rowCount()
        self.check_model.insertRow(new_row)
        self.check_model.setItem(new_row, 0, QStandardItem(item_number))  # №
        self.check_model.setItem(new_row, 1, QStandardItem(item_name))  # Наименование
        self.check_model.setItem(new_row, 2, QStandardItem(item_anime))  # Аниме

        # Удаление строки из каталога
        self.catalog_model.removeRow(source_index.row())

        self.add_delegate_check(self.check_model, self.ui_main_window.chek_table, 3)  # Добавление кнопок увеличения и уменьшения кол-ва товаров в таблицу

        self.ui_main_window.catalog_table.clearSelection()
        self.ui_main_window.confirm_button.setEnabled(True)

        self.total_cost += int(item_price)
        self.ui_main_window.total_label.setText(str(self.total_cost) + " руб.") # Увеличение суммы покупки

    # Функция исключения позиций из чека
    def remove_item_from_check(self):
        # Получаем выделенные индексы в таблице чека
        selected_indexes = self.ui_main_window.chek_table.selectedIndexes()

        if not selected_indexes:
            return  # Если ничего не выбрано

        # Получение индекса строки, выбранной пользователем
        selected_row = selected_indexes[0].row()

        # Получение данных из выбранной строки
        item_number = self.check_model.item(selected_row, 0).text()  # №

        # Извлечение сохраненную информацию о товаре
        item_info = self.item_info.get(item_number)
        if item_info is None:
            return  # Если информация не найдена

        item_name = item_info['name']
        item_anime = item_info['anime']
        item_price = item_info['price']
        item_quantity = item_info['quantity']
        item_select_quantity = item_info['select_quantity']

        # Добавление элемента обратно в каталог
        new_row = self.catalog_model.rowCount()
        self.catalog_model.insertRow(new_row)
        self.catalog_model.setItem(new_row, 0, QStandardItem(item_number))  # №
        self.catalog_model.setItem(new_row, 1, QStandardItem(item_name))  # Наименование
        self.catalog_model.setItem(new_row, 2, QStandardItem(item_anime))  # Аниме
        self.catalog_model.setItem(new_row, 3, QStandardItem(item_price))  # Цену
        self.catalog_model.setItem(new_row, 4, QStandardItem(item_quantity))  # Кол-во

        # Удаление строки из чека
        self.check_model.removeRow(selected_row)

        # Удаление информации о товаре из словаря
        del self.item_info[item_number]
        if not self.item_info:
            self.ui_main_window.confirm_button.setEnabled(False)

        # Очистка выделения в таблице чека
        self.ui_main_window.chek_table.clearSelection()

        self.total_cost -= int(item_price) * int(item_select_quantity)
        self.ui_main_window.total_label.setText(str(self.total_cost) + " руб.") # Уменьшение суммы покупки

    # Создание файла чека/продажа товара
    def save_check(self):
        # Создание нового документа
        doc = Document()

        # Заголовок
        heading = doc.add_heading('Чек', level=1)
        heading.alignment = 1

        # Дата покупки
        purchase_date = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
        doc.add_paragraph(f'Покупка совершена {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.')

        # Общая сумма
        total_sum = 0

        # Заголовки таблицы
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Наименование товара'
        hdr_cells[1].text = 'Количество'
        hdr_cells[2].text = 'Цена за единицу'
        hdr_cells[3].text = 'Общая сумма'

        # Установка стиля таблицы
        table.style = 'Table Grid'  # Добавляет линии к таблице

        # Перенос данных из таблицы чека
        for row in range(self.check_model.rowCount()):
            item_number = self.check_model.item(row, 0).text()  # Получение номера товара
            item_info = self.item_info.get(item_number)

            if item_info is None:
                continue  # Если информация о товаре не найден

            item_name = item_info['name']  # Наименование
            item_delegate = self.ui_main_window.chek_table.indexWidget(self.check_model.index(row, 3))
            item_quantity = item_delegate.get_count()  #  количество
            item_price = int(item_info['price'])  # Цена из словаря
            item_total = item_quantity * item_price  # Общая сумма для товара
            total_sum += item_total

            add_sales(item_name, item_info['anime'], item_price, item_info['quantity'], item_quantity)

            row_cells = table.add_row().cells
            row_cells[0].text = item_name
            row_cells[1].text = str(item_quantity)
            row_cells[2].text = str(item_price)
            row_cells[3].text = str(item_total)

        # Добавление итоговой суммы в таблицу
        total_row_cells = table.add_row().cells
        total_row_cells[0].text = 'Итого'
        total_row_cells[1].text = ''
        total_row_cells[2].text = ''
        total_row_cells[3].text = str(total_sum)

        # Сохранение документа в папку «Чеки»
        check_folder = 'Чеки'
        os.makedirs(check_folder, exist_ok=True)  # Создание папки, если она не существует
        docx_file_path = os.path.join(check_folder, f'чек_{purchase_date}.docx')
        doc.save(docx_file_path)

        # Конвертация в PDF
        pdf_file_path = os.path.join(check_folder, f'чек_{purchase_date}.pdf')
        convert(docx_file_path, pdf_file_path)

        # Удаление .docx файла после конвертации
        os.remove(docx_file_path)

        self.set_model_catalog()
        self.item_info.clear()
        self.check_model.clear()
        self.set_model_check()

    # Открытие окна для редактирования аниме
    def update_create_edit_anime(self, press_edit, current_name):
        self.create_edit_anime.press_edit = press_edit
        self.create_edit_anime.current_name = current_name

    # Открытие окна для редактирования товара
    def update_create_edit_product(self, press_edit, current_name, current_anime, current_price, current_count):
        self.create_edit_product.press_edit = press_edit
        self.create_edit_product.current_name = current_name
        self.create_edit_product.current_anime = current_anime
        self.create_edit_product.current_price = current_price
        self.create_edit_product.current_count = current_count

    # Переключение м/у вкладками
    def change_widget(self, index):
        if index == 0: # Открытие Каталога
            self.ui_main_window.catalog_button.setEnabled(False)
            self.ui_main_window.sales_button.setEnabled(True)
            self.ui_main_window.management_button.setEnabled(True)
            self.ui_main_window.confirm_button.setEnabled(False)

            self.ui_main_window.search_field.clear()
            self.set_model_catalog()

            self.check_model.clear() # Очистка таблицы чек
            self.set_model_check()

            self.total_cost = 0
            self.ui_main_window.total_label.setText(str(self.total_cost) + " руб.")

        if index == 1: # Открытие Продажи
            self.ui_main_window.catalog_button.setEnabled(True)
            self.ui_main_window.sales_button.setEnabled(False)
            self.ui_main_window.management_button.setEnabled(True)

            self.ui_main_window.begin_dateEdit.setDate(datetime.now())
            self.ui_main_window.end_dateEdit.setDate(datetime.now())
            self.set_model_sale()

        if index == 2: # Открытие Управление
            self.ui_main_window.catalog_button.setEnabled(True)
            self.ui_main_window.sales_button.setEnabled(True)
            self.ui_main_window.management_button.setEnabled(False)

        self.ui_main_window.stacked_widget.setCurrentIndex(index)

    # Открытие окна для добавления товара/аниме
    def open_create(self):
        if self.ui_main_window.productstable_button.isEnabled(): # Если активна кнопка "Товары", то открыта таблица Аниме
            self.create_edit_anime.ui_create_edit_anime.name_textedit.clear()
            self.create_edit_anime.setWindowModality(Qt.ApplicationModal)
            self.create_edit_anime.show()
        else:
            self.create_edit_product.clear_text()
            self.create_edit_product.setWindowModality(Qt.ApplicationModal)
            self.create_edit_product.show()

    # Открытие окна для отправки на печать
    def open_preview(self):
        # Создаем экземпляр Preview и передаем даты
        self.preview_window = Preview(self.sale_model, self.proxy_model_sale)

        self.preview_window.show()


    # Добавление кнопок в столбик действие в таблице "Товары"
    def add_delegate_management(self, model, table, number):
        for row in range(model.rowCount()):
            item = model.item(row, 1)

            if item is not None:
                item_delegate = ItemDelegateData(row)  # экземпляр ItemDelegateData
                table.setIndexWidget(model.index(row, number), item_delegate)  # Установка кнопок редактирования и удаления в таблицу
                item_delegate.button_clicked.connect(self.handle_button_click_management)

    # Функция для кнопок в таблице "Товары"
    def handle_button_click_management(self, button_type, row_number):
        type_button = button_type # Определение кнопки
        number_row = row_number # Определение в какой она строке

        if type_button == "edit":
            if self.ui_main_window.productstable_button.isEnabled():
                index_name = self.ui_main_window.management_table.model().index(number_row, 1)

                # Извлечение текста из ячеек
                self.create_edit_anime.ui_create_edit_anime.name_textedit.setText(index_name.data())

                current_name = index_name.data()  #Запоминание текста из ячеек

                self.create_edit_anime.show()

                press_edit = True
                self.update_create_edit_anime(press_edit, current_name)
            else:
                # Получение индекса
                index_name = self.ui_main_window.management_table.model().index(number_row, 1)
                index_anime = self.ui_main_window.management_table.model().index(number_row, 2)
                index_price = self.ui_main_window.management_table.model().index(number_row, 3)
                index_count = self.ui_main_window.management_table.model().index(number_row, 4)

                # Извлечение текста из ячеек
                self.create_edit_product.ui_create_edit_product.name_textedit.setText(index_name.data())
                self.create_edit_product.ui_create_edit_product.anime_textedit.setText(index_anime.data())
                self.create_edit_product.ui_create_edit_product.price_textedit.setText(index_price.data())
                self.create_edit_product.ui_create_edit_product.count_textedit.setText(index_count.data())

                current_name = index_name.data() #Запоминание текста из ячеек
                current_anime = index_anime.data()
                current_price = index_price.data()
                current_count = index_count.data()

                self.create_edit_product.show()

                press_edit = True
                self.update_create_edit_product(press_edit, current_name, current_anime, current_price,
                                                current_count)
        else:
            if self.ui_main_window.productstable_button.isEnabled():
                index_name = self.ui_main_window.management_table.model().index(number_row, 1)

                anime_name = index_name.data()
                delete_anime_by_name(anime_name)

                # Загрузка данных из базы данных
                load_data_from_db(self.anime_model, 'anime_table')
                self.anime_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Действие'])
                self.add_delegate_management(self.anime_model, self.ui_main_window.management_table, 2)

            else:
                index_name = self.ui_main_window.management_table.model().index(number_row, 1)
                index_anime = self.ui_main_window.management_table.model().index(number_row, 2)
                index_price = self.ui_main_window.management_table.model().index(number_row, 3)
                index_count = self.ui_main_window.management_table.model().index(number_row, 4)

                # Извлечение текста из ячеек
                name = index_name.data()
                anime = index_anime.data()
                price = index_price.data()
                count = index_count.data()
                delete_product(name, anime, price, count)

                load_data_from_db(self.product_model, 'product_table')
                self.product_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Аниме', 'Цена', 'Кол-во', 'Действие'])
                self.add_delegate_management(self.product_model, self.ui_main_window.management_table, 5)

    # Функция сохраняющая БД в sql файл после закрытия программы
    def closeEvent(self, event):
        save_database()
        event.accept()

class CreateEditProduct(QtWidgets.QWidget):
    def __init__(self, main_window, product_model, add_delegate):
        super(CreateEditProduct, self).__init__()
        self.ui_create_edit_product = Ui_CreateEditProduct_form()
        self.ui_create_edit_product.setupUi(self)

        self.product_model = product_model
        self.main_window = main_window
        self.add_delegate = add_delegate

        self.press_edit = False
        self.current_name = None
        self.current_anime = None
        self.current_price = None
        self.current_count = None

        self.ui_create_edit_product.ok_button.clicked.connect(self.save_product) # Кнопка подтверждения
        self.ui_create_edit_product.cancel_button.clicked.connect(self.close_create_edit_product) # Кнопка отмены

    def save_product(self):
        name = self.ui_create_edit_product.name_textedit.toPlainText().strip()  # Наименование
        anime = self.ui_create_edit_product.anime_textedit.toPlainText().strip()  # Аниме
        price = self.ui_create_edit_product.price_textedit.toPlainText().strip()  # Цена
        count = self.ui_create_edit_product.count_textedit.toPlainText().strip()  # Количество

        def is_valid_input(value):
            # Проверка, что строка состоит только из букв и цифр
            return bool(re.match(r'^[\w\s]+$', value))  # \w соответствует буквам и цифрам, \s соответствует пробелам

        def is_valid_number(number):
            # Проверка, что count состоит только из цифр
            return number.isdigit()

        if self.press_edit:
            # Проверка входных данных
            input_valid = (
                    is_valid_input(name) and
                    is_valid_input(anime) and
                    is_valid_number(price) and
                    is_valid_number(count)
            )

            if not input_valid:
                # Очистка полей, если входные данные недействительны
                if not is_valid_number(price):
                    self.ui_create_edit_product.price_textedit.setText(self.current_price) # Ввод текущей цены
                if not is_valid_number(count):
                    self.ui_create_edit_product.count_textedit.setText(self.current_count)
                if not is_valid_input(anime):
                    self.ui_create_edit_product.anime_textedit.setText(self.current_anime)
                if not is_valid_input(name):
                    self.ui_create_edit_product.name_textedit.setText(self.current_name)
                return

            # Если входные данные валидны, обновляем продукт
            if update_product(self.current_name, self.current_anime, self.current_price, self.current_count, name,
                              anime, price, count):
                self.close()
                self.press_edit = False

        else:
            print('New')
            # Проверка входных данных и добавление новой записи
            if not is_valid_input(name) or not is_valid_input(anime) or not is_valid_number(price) or not is_valid_number(
                    count) or not add_to_database_product(name, anime, price, count):
                # Очистка полей и уведомление пользователя об ошибке
                if not is_valid_input(name):
                    self.ui_create_edit_product.name_textedit.clear()
                if not is_valid_number(price):
                    self.ui_create_edit_product.price_textedit.clear()  # Очистка поля цены
                if not is_valid_number(count):
                    self.ui_create_edit_product.count_textedit.clear()  # Очистка поля количества
                if not is_valid_input(anime) or not add_to_database_product(name, anime, price, count):
                    self.ui_create_edit_product.anime_textedit.clear()  # Очистка поля аниме
                return

        self.clear_text()# Очистка полей

        # Загрузка данных из базы данных
        load_data_from_db(self.product_model, 'product_table')
        self.product_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Аниме', 'Цена', 'Кол-во', 'Действие'])
        self.add_delegate(self.product_model, self.main_window.management_table, 5)

    def close_create_edit_product(self):
        self.clear_text()
        self.close()

    def clear_text(self):
        self.ui_create_edit_product.name_textedit.clear()
        self.ui_create_edit_product.anime_textedit.clear()
        self.ui_create_edit_product.price_textedit.clear()
        self.ui_create_edit_product.count_textedit.clear()

class CreateEditAnime(QtWidgets.QWidget):
    def __init__(self, main_window, anime_model, add_delegate):
        super(CreateEditAnime, self).__init__()
        self.ui_create_edit_anime = Ui_CreateEditAnime_form()
        self.ui_create_edit_anime.setupUi(self)

        self.main_window = main_window
        self.anime_model = anime_model
        self.add_delegate = add_delegate

        self.press_edit = False
        self.current_name = None

        self.ui_create_edit_anime.ok_button.clicked.connect(self.save_anime) # Кнопка подтверждения
        self.ui_create_edit_anime.cancel_button.clicked.connect(self.close_create_edit_anime) # Кнопка отмены

    def save_anime(self):
        name = self.ui_create_edit_anime.name_textedit.toPlainText().strip()  # Наименование

        def is_valid_input(value):
            # Проверка, что строка состоит только из букв и цифр
            return bool(re.match(r'^[\w\s]+$', value))

        # Проверка входных данных
        if not is_valid_input(name) :
            self.ui_create_edit_anime.name_textedit.clear()
            return

        if self.press_edit:
            update_anime(self.current_name, name)
            self.close()
            self.press_edit = False
        else:
            add_to_database_anime(name) # Добавление новой записи

        self.ui_create_edit_anime.name_textedit.clear()   # Очистка поля

        # Загрузка данных из базы данных
        load_data_from_db(self.anime_model, 'anime_table')
        self.anime_model.setHorizontalHeaderLabels(['№', 'Наименование', 'Действие'])
        self.add_delegate(self.anime_model, self.main_window.management_table, 2)

    def close_create_edit_anime(self):
        self.ui_create_edit_anime.name_textedit.clear() # Очистка поля
        self.close()


class Preview(QtWidgets.QWidget):
    def __init__(self, model, filter_model):
        super(Preview, self).__init__()
        self.ui_preview = Ui_preview_form()
        self.ui_preview.setupUi(self)

        self.sale_model = model
        self.filter_model = filter_model

        self.start_date, self.end_date = self.get_sales_date_range()

        self.ui_preview.save_button.clicked.connect(self.save_report)

        # Отображение pdf файла
        self.pdf_file_path = self.create_sales_report()

        self.ui_preview.viewer.load(QtCore.QUrl.fromLocalFile(os.path.abspath(self.pdf_file_path)))

    def get_sales_date_range(self):
        # Переменные для минимальной и максимальной даты периода
        min_date = None
        max_date = None

        # Перенос данных из отфильтрованной модели продаж
        for row in range(self.filter_model.rowCount()):
            source_row = self.filter_model.mapToSource(self.filter_model.index(row, 0)).row()
            sale_date = self.sale_model.item(source_row, 5).text()
            sale_date = sale_date.split()[0]
            sale_date = datetime.strptime(sale_date, '%Y-%m-%d')  # Преобразование строки в дату

            if min_date is None or sale_date < min_date:
                min_date = sale_date
            if max_date is None or sale_date > max_date:
                max_date = sale_date

        return min_date, max_date

    def create_sales_report(self):
        # Создание нового документа
        doc = Document()

        # Заголовок
        heading = doc.add_heading('Отчет о продажах', level=1)
        heading.alignment = 1

        # Введение
        doc.add_paragraph('Цель данного отчета - предоставить обзор продаж за указанный период.')
        doc.add_paragraph('Период отчета: с {} по {}.'.format(
            self.start_date.strftime('%Y-%m-%d'),
            self.end_date.strftime('%Y-%m-%d')
        ))

        total_quantity = 0
        total_sum = 0
        item_sales = {}

        # Перенос данных из отфильтрованной модели продаж
        for row in range(self.filter_model.rowCount()):
            source_row = self.filter_model.mapToSource(self.filter_model.index(row, 0)).row()
            item_name = self.sale_model.item(source_row, 1).text()  # Наименование товара
            anime_name = self.sale_model.item(source_row, 2).text()  # Наименование аниме
            item_price = int(self.sale_model.item(source_row, 3).text())  # Цена
            item_quantity = int(self.sale_model.item(source_row, 4).text())  # Количество

            # Дата продажи
            sale_date = self.sale_model.item(source_row, 5).text()
            sale_date = sale_date.split()[0]
            sale_date = datetime.strptime(sale_date, '%Y-%m-%d')  # Преобразование строки в дату

            # Попадает ли дата продажи в отчетный период
            if self.start_date <= sale_date <= self.end_date:
                total_quantity += item_quantity
                item_total = item_quantity * item_price
                total_sum += item_total

                # Объединение наименования товара и названия аниме
                combined_name = f"{item_name} {anime_name}"

                # Сохранение информации о продажах по товарам, суммируя количество и общую стоимость
                if combined_name in item_sales:
                    item_sales[combined_name]['quantity'] += item_quantity
                    item_sales[combined_name]['total'] += item_total
                else:
                    item_sales[combined_name] = {
                        'quantity': item_quantity,
                        'total': item_total,
                        'price': item_price
                    }

        # Добавление общих данных о продажах
        doc.add_paragraph(f'Общее количество проданных товаров: {total_quantity}.')
        doc.add_paragraph(f'Общая сумма продажи: {total_sum}.')

        # Наиболее продаваемый товар
        if item_sales:
            best_selling_item = max(item_sales.items(), key=lambda x: x[1]['quantity'])
            best_item_name = best_selling_item[0]
            best_item_quantity = best_selling_item[1]['quantity']
            best_item_share = (best_item_quantity / total_quantity * 100) if total_quantity > 0 else 0

            doc.add_paragraph('Наиболее продаваемый товар:')
            doc.add_paragraph(f'1. Наименование: {best_item_name};')
            doc.add_paragraph(f'2. Количество проданного: {best_item_quantity};')
            doc.add_paragraph(f'3. Доля в общем объеме продаж: {best_item_share:}%.')

        # Аниме, товары с которым пользовались наибольшим спросом
        anime_sales = {}
        for row in range(self.filter_model.rowCount()):
            source_row = self.filter_model.mapToSource(self.filter_model.index(row, 0)).row()
            anime_name = self.sale_model.item(source_row, 2).text()  # Название аниме
            item_quantity = int(self.sale_model.item(source_row, 4).text())  # Количество

            if anime_name not in anime_sales:
                anime_sales[anime_name] = 0
            anime_sales[anime_name] += item_quantity

        if anime_sales:
            most_popular_anime = max(anime_sales.items(), key=lambda x: x[1])
            anime_name = most_popular_anime[0]
            anime_quantity = most_popular_anime[1]
            anime_share = (anime_quantity / total_quantity * 100) if total_quantity > 0 else 0

            doc.add_paragraph('Аниме, товары с которым пользовались наибольшим спросом:')
            doc.add_paragraph(f'1. Наименование: {anime_name};')
            doc.add_paragraph(f'2. Количество проданного: {anime_quantity};')
            doc.add_paragraph(f'3. Доля в общем объеме продаж: {anime_share}%.')

        # Продажи по товарам
        doc.add_paragraph('Продажи по товарам:')
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Наименование товара'
        hdr_cells[1].text = 'Количество проданного'
        hdr_cells[2].text = 'Цена за единицу'
        hdr_cells[3].text = 'Общая сумма'

        # Установка стиля таблицы
        table.style = 'Table Grid'

        # Заполнение таблицы данными о продажах
        for item_name, sales_info in item_sales.items():
            row_cells = table.add_row().cells
            row_cells[0].text = item_name  # Объединенное наименование товара и аниме
            row_cells[1].text = str(sales_info['quantity'])  # Количество
            row_cells[2].text = str(sales_info['price'])  # Цена
            row_cells[3].text = str(sales_info['total'])  # Общая

        # Добавление итоговой суммы в таблицу
        total_row_cells = table.add_row().cells
        total_row_cells[0].text = 'Итого'
        total_row_cells[1].text = str(total_quantity)  # Количество
        total_row_cells[2].text = ''
        total_row_cells[3].text = str(total_sum)  # Общая сумма

        # Cоздание пути для файла
        start_date = str(self.start_date).split()[0]
        end_date = str(self.end_date).split()[0]
        report_file_path = f"resources/отчёт_по_продажам_с_{start_date}_по_{end_date}.docx"

        if report_file_path:
            doc.save(report_file_path)

            # Конвертация в PDF
            pdf_file_path = report_file_path.replace('.docx', '.pdf')
            convert(report_file_path, pdf_file_path)

            # Удаление .docx файла после конвертации
            if os.path.exists(report_file_path):
                os.remove(report_file_path)

            return pdf_file_path

    def save_report(self):
        # Сохранение существующего PDF файла в новое место
        options = QFileDialog.Options()
        report_file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить отчет о продажах",
                                                          self.pdf_file_path, "PDF файлы (*.pdf);;Все файлы (*)",
                                                          options=options)

        if report_file_path:
            # Копирование файла по новому пути
            os.rename(self.pdf_file_path, report_file_path)
            self.close()


if __name__ == "__main__":
    app = QtWidgets.QApplication([])
    application = MainWindow()
    application.show()

    sys.exit(app.exec())