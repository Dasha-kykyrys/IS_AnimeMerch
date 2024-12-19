from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH

# создание пустого документа
doc = Document()
# данные таблицы без названий колонок
items = (
    (7, '1024', 'Плюшевые котята'),
    (3, '2042', 'Меховые пчелы'),
    (1, '1288', 'Ошейники для пуделей'),
    (1, '1288', 'Ошейники для пуделей'),
)
# добавляем таблицу с одной строкой
# для заполнения названий колонок
table = doc.add_table(1, len(items[0]))
# определяем стиль таблицы
table.style = 'Light Shading Accent 1'
# Получаем строку с колонками из добавленной таблицы
head_cells = table.rows[0].cells
# добавляем названия колонок
for i, item in enumerate(['Кол-во', 'ID', 'Описание']):
    p = head_cells[i].paragraphs[0]
    # название колонки
    p.add_run(item).bold = True
    # выравниваем посередине
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
# добавляем данные к существующей таблице
for row in items:
    # добавляем строку с ячейками к объекту таблицы
    cells = table.add_row().cells
    for i, item in enumerate(row):
        # вставляем данные в ячейки
        cells[i].text = str(item)
        # если последняя ячейка
        if i == 2:
            # изменим шрифт
            cells[i].paragraphs[0].runs[0].font.name = 'Arial'
doc.save('test.docx')